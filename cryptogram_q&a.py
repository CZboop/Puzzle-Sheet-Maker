import random
import pandas as pd
import operator
import docx
from docx import Document
from docx.shared import Cm
import random
from collections import Counter
import datetime
from docx.shared import Pt

# put phrases to convert here
phrase_list = []

#main function does everything
# has some default arguments that can also be adjusted:
# - save_append is what the end of the filename will be, defaults to a random number
# - type can be either 'mixed' for a mixture of caesar and substitution ciphers, 'sub' for just substitution, 'shift' or any other string for caesar only
# - title will be used within headings in the file
def make_cryptograms(input, save_append=random.randrange(1000), type='mixed', title='Cryptograms'):
    # shuffling phrases that will be used to make cryptograms
    random.shuffle(input)

    #creating a docx document to write the questions to
    document = Document()

    # creating a separate docx document to write the answers to
    answer_document = Document()

    # adding the title provided
    document.add_heading(title.upper())

    # adding a heading to the answer document
    answer_document.add_heading('ANSWERS - ' + title)
    answer_document.add_paragraph()

    # main loop goes through each word or phrase in the input list to create the cryptogram and add to both documents
    # using enumerate so each question can be numbered
    for count,i in enumerate(input):
        # adding number and spacing
        answer_document.add_paragraph(str(count + 1) + ".")
        answer_document.add_paragraph(i)

        # getting a number to decide which type of cipher it will be, may be random or not depending on type argument
        random_for_type = random.randrange(1,10) if type=='mixed' else 10 if type=='sub' else 1

        # setup for creating the puzzle
        alphabet = "abcdefghijklmnopqrstuvwxyz"
        keys = {}
        cryptogram = ""

        # settiing up subsitution with shuffled alphabet assigned to keys dictionary
        if random_for_type>5:
            alphabet_list = list(alphabet)
            random.shuffle(alphabet_list)
            shuffled_alphabet = ''.join(alphabet_list)
            for c,v in enumerate(alphabet):
                keys[v] = shuffled_alphabet[c]

        # making shifted/caesar cipher, picks a number to shift by and uses ascii values to shift by this amount
        else:
            shift_num_options = [i for i in range(-25,25) if i!=0]
            shift_num = random.choice(shift_num_options)
            for c,v in enumerate(alphabet):
                keys[v] = chr((ord(v) + shift_num) % 26 + 97)
        for letter in i.lower():
            cryptogram += keys.get(letter, letter)

        # adding number and question to question document
        document.add_paragraph(str(count + 1) + ".")
        document.add_paragraph(cryptogram.upper())

        # adding underscores to 'fill in' for each letter, preserving spaces and other characters
        underscores = ""
        for letter in cryptogram:
            if letter.isalpha():
                underscores += "_"
            else:
                underscores += letter

        # adding the underscore string to the question document
        underscores_para = document.add_paragraph(underscores)

        document.add_paragraph()

        # add grid/table for each letter, with alphabet on top row and empty cell underneath
        document.add_paragraph("Letter Key:")
        letters_table1 = document.add_table(rows=2, cols=13)
        letters_table1.style = 'Table Grid'
        for x,y in enumerate(alphabet[:13]):
            letters_table1.cell(0,x).text = y.upper()
        for row in letters_table1.rows:
            row.height = Cm(0.7)

        document.add_paragraph()

        # split into first and second half of alphabet so big enough be written in
        letters_table2 = document.add_table(rows=2, cols=13)
        letters_table2.style = 'Table Grid'
        for x,y in enumerate(alphabet[13:]):
            letters_table2.cell(0,x).text = y.upper()
        for row in letters_table2.rows:
            row.height = Cm(0.7)

        # adding solutions in terms of what letters match into the answer document
        answer_document.add_paragraph("Letter Key: Solution")
        letters_table_ans1 = answer_document.add_table(rows=2, cols=13)
        letters_table_ans1.style = 'Table Grid'
        for x,y in enumerate(alphabet[:13]):
            letters_table_ans1.cell(0,x).text = y.upper()
            letters_table_ans1.cell(1,x).text = keys.get(y).upper()
        for row in letters_table_ans1.rows:
            row.height = Cm(0.7)

        answer_document.add_paragraph()

        # again in two tables
        letters_table_ans2 = answer_document.add_table(rows=2, cols=13)
        letters_table_ans2.style = 'Table Grid'
        for x,y in enumerate(alphabet[13:]):
            letters_table_ans2.cell(0,x).text = y.upper()
            letters_table_ans2.cell(1,x).text = keys.get(y).upper()
        for row in letters_table_ans2.rows:
            row.height = Cm(0.7)

        answer_document.add_paragraph()

        document.add_paragraph()

        # adding letter frequencies, using built in counter and adding these in another alphabet table
        document.add_paragraph("Letter Frequencies:")
        counts = Counter(cryptogram)
        letter_freq_table1 = document.add_table(rows=2, cols=13)
        letter_freq_table1.style = 'Table Grid'
        for x,y in enumerate(alphabet[:13]):
            letter_freq_table1.cell(0,x).text = y.upper()
            letter_freq_table1.cell(1,x).text = str(counts.get(y, 0))
        for row in letter_freq_table1.rows:
            row.height = Cm(0.7)

        document.add_paragraph()

        letter_freq_table2 = document.add_table(rows=2, cols=13)
        letter_freq_table2.style = 'Table Grid'
        for x,y in enumerate(alphabet[13:]):
            letter_freq_table2.cell(0,x).text = y.upper()
            letter_freq_table2.cell(1,x).text = str(counts.get(y, 0))
        for row in letter_freq_table2.rows:
            row.height = Cm(0.7)

        document.add_paragraph()

    # saving each document using either default random number or argument given when calling the function
    document.save("cryptograms_{0}.docx".format(save_append))
    answer_document.save("cryptograms_answers_{0}.docx".format(save_append))

# running with defaults and the list from the start of the file
make_cryptograms(input)
